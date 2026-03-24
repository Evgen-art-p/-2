# studio/workshop_export.py — Экспорт в Word / PDF
# Вынесено из ui_workshop.py (строки 68-871)

import re


def _get_viewer_text(state):
    """Собирает текст: viewer > brief > runs файлы"""
    # 1. Из viewer
    v = state.get("viewer_content", "")
    if v and v.strip():
        return v
    # 2. Из брифа
    b = state.get("master_brief", "")
    if b and b.strip():
        return b
    # 3. Из результатов агентов
    for wid in reversed(list(state.get("results", {}).keys())):
        res = state["results"][wid]
        if isinstance(res, dict):
            t = res.get("text", "")
        else:
            t = str(res)
        if t and t.strip():
            return t
    # 4. Ищем ВСЕ .md файлы в runs/ рекурсивно
    from pathlib import Path as _PP
    runs_dir = _PP("runs")
    if runs_dir.exists():
        all_md = sorted(runs_dir.rglob("*.md"), key=lambda f: f.stat().st_mtime, reverse=True)
        for f in all_md:
            if f.name == "README.md":
                continue
            try:
                txt = f.read_text(encoding="utf-8")
                if txt.strip() and len(txt) > 50:
                    print(f"[EXPORT] Нашёл: {f}")
                    return txt
            except:
                pass
    print("[EXPORT] Ничего не нашёл!")
    return ""


def _export_docx(text, filepath):
    """Генерирует .docx с заголовками, списками и таблицами"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        import re as _re

        doc = Document()

        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)

        def clean_md(s):
            s = _re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', s)
            s = _re.sub(r'\*\*(.+?)\*\*', r'\1', s)
            s = _re.sub(r'\*(.+?)\*', r'\1', s)
            s = _re.sub(r'`(.+?)`', r'\1', s)
            return s.strip()

        def add_table(table_lines):
            rows = []
            for tl in table_lines:
                tl = tl.strip()
                if tl.startswith("|"):
                    tl = tl[1:]
                if tl.endswith("|"):
                    tl = tl[:-1]
                cells = [clean_md(c.strip()) for c in tl.split("|")]
                if all(_re.match(r'^[:\-\s]+$', c) for c in cells if c):
                    continue
                if any(c.strip() for c in cells):
                    rows.append(cells)

            if len(rows) < 2:
                return

            ncols = max(len(r) for r in rows)
            for r in rows:
                while len(r) < ncols:
                    r.append("")

            table = doc.add_table(rows=len(rows), cols=ncols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT

            for ri, row in enumerate(rows):
                for ci, cell_text in enumerate(row):
                    cell = table.cell(ri, ci)
                    cell.text = cell_text

                    for paragraph in cell.paragraphs:
                        paragraph.style.font.size = Pt(9)
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
                            run.font.name = "Arial"

                    if ri == 0:
                        shading = parse_xml(
                            f'<w:shd {nsdecls("w")} w:fill="2D3748"/>'
                        )
                        cell._tc.get_or_add_tcPr().append(shading)
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.color.rgb = RGBColor(255, 255, 255)
                                run.font.bold = True

            doc.add_paragraph("")

        all_lines = text.split("\n")
        idx = 0

        while idx < len(all_lines):
            line = all_lines[idx]
            stripped = line.strip()

            if not stripped:
                doc.add_paragraph("")
                idx += 1
                continue

            if "|" in stripped and stripped.count("|") >= 2:
                tbuf = [stripped]
                j = idx + 1
                while j < len(all_lines):
                    nl = all_lines[j].strip()
                    if "|" in nl and nl.count("|") >= 2:
                        tbuf.append(nl)
                        j += 1
                    elif nl and all(c in "-|: " for c in nl):
                        tbuf.append(nl)
                        j += 1
                    else:
                        break
                if len(tbuf) >= 2:
                    add_table(tbuf)
                    idx = j
                    continue

            clean = clean_md(stripped)

            if stripped.startswith("---"):
                idx += 1
                continue

            if not clean:
                idx += 1
                continue

            if stripped.startswith("#### "):
                doc.add_heading(clean_md(stripped[5:]), level=4)
            elif stripped.startswith("### "):
                doc.add_heading(clean_md(stripped[4:]), level=3)
            elif stripped.startswith("## "):
                doc.add_heading(clean_md(stripped[3:]), level=2)
            elif stripped.startswith("# "):
                doc.add_heading(clean_md(stripped[2:]), level=1)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(clean_md(stripped[2:]), style="List Bullet")
            elif _re.match(r'^\d+\.\s', stripped):
                txt = _re.sub(r'^\d+\.\s*', '', stripped)
                doc.add_paragraph(clean_md(txt), style="List Number")
            else:
                doc.add_paragraph(clean)

            idx += 1

        doc.save(str(filepath))
        return True
    except ImportError:
        return False
    except Exception as ex:
        print(f"[DOCX ERROR] {ex}")
        import traceback
        traceback.print_exc()
        return False


def _export_pdf(text, filepath):
    """PDF с кириллицей, таблицами landscape и переносом текста"""
    try:
        from fpdf import FPDF
        import os, re

        class MyPDF(FPDF):
            def __init__(self):
                super().__init__()
                self._font_name = "Helvetica"
                self._has_bold = False

            def setup_font(self):
                font_paths = [
                    "C:/Windows/Fonts/arial.ttf",
                    "C:/Windows/Fonts/calibri.ttf",
                    "C:/Windows/Fonts/tahoma.ttf",
                ]
                for fp in font_paths:
                    if os.path.exists(fp):
                        try:
                            fname = os.path.splitext(os.path.basename(fp))[0]
                            self.add_font(fname, "", fp, uni=True)
                            self._font_name = fname
                            bd = fp.replace(".ttf", "bd.ttf")
                            if os.path.exists(bd):
                                self.add_font(fname, "B", bd, uni=True)
                                self._has_bold = True
                            return
                        except:
                            continue

            def set_f(self, size=10, bold=False):
                if bold and self._has_bold:
                    self.set_font(self._font_name, "B", size)
                else:
                    self.set_font(self._font_name, "", size)

            def cleaned(self, s):
                s = re.sub(r"\*\*\*(.+?)\*\*\*", r"\1", s)
                s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
                s = re.sub(r"\*(.+?)\*", r"\1", s)
                s = re.sub(r"`(.+?)`", r"\1", s)
                return s.strip()

            def get_text_lines(self, text, width):
                """Разбивает текст на строки по ширине ячейки"""
                if not text:
                    return [""]
                words = text.split(" ")
                result_lines = []
                current = ""
                for word in words:
                    test = current + (" " if current else "") + word
                    if self.get_string_width(test) <= width - 2:
                        current = test
                    else:
                        if current:
                            result_lines.append(current)
                        current = word
                if current:
                    result_lines.append(current)
                return result_lines if result_lines else [""]

            def render_table(self, table_lines):
                """Таблица с переносом текста и авто-высотой строк"""
                rows = []
                for tl in table_lines:
                    tl = tl.strip()
                    if tl.startswith("|"):
                        tl = tl[1:]
                    if tl.endswith("|"):
                        tl = tl[:-1]
                    cells = [self.cleaned(c.strip()) for c in tl.split("|")]
                    if all(re.match(r'^[:\-\s]+$', c) for c in cells if c):
                        continue
                    if any(c.strip() for c in cells):
                        rows.append(cells)

                if len(rows) < 2:
                    return

                ncols = max(len(r) for r in rows)
                for r in rows:
                    while len(r) < ncols:
                        r.append("")

                is_wide = ncols >= 5
                if is_wide:
                    self.add_page(orientation="L")
                    page_w = self.w - self.l_margin - self.r_margin
                    self.set_f(8)
                else:
                    page_w = self.w - self.l_margin - self.r_margin
                    self.set_f(9)

                col_w = []
                for ci in range(ncols):
                    max_w = 12
                    for r in rows:
                        sw = self.get_string_width(r[ci][:60]) + 6
                        max_w = max(max_w, sw)
                    col_w.append(max_w)

                total = sum(col_w)
                if total > page_w:
                    ratio = page_w / total
                    col_w = [max(c * ratio, 10) for c in col_w]

                total2 = sum(col_w)
                if total2 < page_w:
                    col_w[-1] += page_w - total2

                cell_h = 5 if is_wide else 6

                for ri, row in enumerate(rows):
                    is_header = (ri == 0)

                    if is_header:
                        self.set_f(8 if is_wide else 9, bold=True)
                    else:
                        self.set_f(8 if is_wide else 9)

                    wrapped = []
                    max_lines = 1
                    for ci, cell in enumerate(row):
                        if ci < len(col_w):
                            cell_lines = self.get_text_lines(cell, col_w[ci])
                            wrapped.append(cell_lines)
                            max_lines = max(max_lines, len(cell_lines))
                        else:
                            wrapped.append([""])

                    row_h = cell_h * max_lines

                    if self.get_y() + row_h > self.h - 15:
                        self.add_page(orientation="L" if is_wide else "P")

                    x_start = self.l_margin
                    y_start = self.get_y()

                    if is_header:
                        self.set_fill_color(45, 55, 72)
                        self.set_text_color(255, 255, 255)
                    else:
                        if ri % 2 == 0:
                            self.set_fill_color(248, 248, 252)
                        else:
                            self.set_fill_color(255, 255, 255)
                        self.set_text_color(30, 30, 30)

                    for ci in range(min(ncols, len(col_w))):
                        x = x_start + sum(col_w[:ci])
                        self.rect(x, y_start, col_w[ci], row_h, "DF")

                        cell_lines = wrapped[ci] if ci < len(wrapped) else [""]
                        for li, cl in enumerate(cell_lines):
                            self.set_xy(x + 1, y_start + li * cell_h)
                            self.cell(col_w[ci] - 2, cell_h, cl, border=0)

                    self.set_xy(x_start, y_start + row_h)

                self.set_text_color(30, 30, 30)
                self.ln(6)

                if is_wide:
                    self.add_page(orientation="P")
                    self.set_f(10)

        pdf = MyPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.setup_font()
        pdf.set_f(10)

        pw = pdf.w - pdf.l_margin - pdf.r_margin
        all_lines = text.split("\n")
        idx = 0

        while idx < len(all_lines):
            line = all_lines[idx]
            stripped = line.strip()

            if not stripped:
                pdf.ln(3)
                idx += 1
                continue

            if "|" in stripped and stripped.count("|") >= 2:
                tbuf = [stripped]
                j = idx + 1
                while j < len(all_lines):
                    nl = all_lines[j].strip()
                    if "|" in nl and nl.count("|") >= 2:
                        tbuf.append(nl)
                        j += 1
                    elif nl and all(c in "-|: " for c in nl):
                        tbuf.append(nl)
                        j += 1
                    else:
                        break
                if len(tbuf) >= 2:
                    pdf.render_table(tbuf)
                    idx = j
                    continue

            clean = pdf.cleaned(stripped)
            if not clean or clean == "---":
                idx += 1
                continue

            if stripped.startswith("#### "):
                pdf.set_f(11, bold=True)
                pdf.multi_cell(pw, 6, pdf.cleaned(stripped[5:]))
                pdf.set_f(10)
                pdf.ln(2)
            elif stripped.startswith("### "):
                pdf.set_f(12, bold=True)
                pdf.multi_cell(pw, 7, pdf.cleaned(stripped[4:]))
                pdf.set_f(10)
                pdf.ln(2)
            elif stripped.startswith("## "):
                pdf.set_f(14, bold=True)
                pdf.multi_cell(pw, 8, pdf.cleaned(stripped[3:]))
                pdf.set_f(10)
                pdf.ln(3)
            elif stripped.startswith("# "):
                pdf.set_f(18, bold=True)
                pdf.multi_cell(pw, 10, pdf.cleaned(stripped[2:]))
                pdf.set_f(10)
                pdf.ln(4)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                pdf.multi_cell(pw, 6, "  \u2022 " + pdf.cleaned(stripped[2:]))
                pdf.ln(1)
            elif re.match(r"^\d+\.\s", stripped):
                num = re.match(r"^(\d+)\.", stripped).group(1)
                txt = re.sub(r"^\d+\.\s*", "", stripped)
                pdf.multi_cell(pw, 6, f"  {num}. " + pdf.cleaned(txt))
                pdf.ln(1)
            else:
                pdf.multi_cell(pw, 6, clean)
                pdf.ln(1)

            idx += 1

        pdf.output(str(filepath))
        return True
    except ImportError:
        return False
    except Exception as ex:
        print(f"[PDF ERROR] {ex}")
        import traceback
        traceback.print_exc()
        return False
